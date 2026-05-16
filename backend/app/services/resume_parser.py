"""Resume parsing service.

Extracts text from PDF/DOCX files and uses Claude to parse structured data.
"""

import json
import logging
import re
import time
from io import BytesIO
from pathlib import Path

import anthropic
import pdfplumber
from pydantic import BaseModel, ValidationError

from app.core.config import settings
from app.core.security_config import security_settings

logger = logging.getLogger(__name__)


# ── Exceptions ─────────────────────────────────────────────────────────────────

class ResumeParseError(Exception):
    """Raised when resume parsing fails after all retries."""


class TextExtractionError(Exception):
    """Raised when text cannot be extracted from a file."""


# ── Pydantic schemas ────────────────────────────────────────────────────────────

class WorkExperience(BaseModel):
    company: str
    title: str
    start_date: str | None = None
    end_date: str | None = None
    location: str | None = None
    description: list[str] = []


class Education(BaseModel):
    institution: str
    degree: str | None = None
    field: str | None = None
    graduation_year: str | None = None
    gpa: str | None = None


class Project(BaseModel):
    name: str
    description: str | None = None
    technologies: list[str] = []
    url: str | None = None


class ResumeData(BaseModel):
    full_name: str | None = None
    email: str | None = None
    phone: str | None = None
    location: str | None = None
    linkedin_url: str | None = None
    github_url: str | None = None
    portfolio_url: str | None = None
    summary: str | None = None
    work_experience: list[WorkExperience] = []
    education: list[Education] = []
    skills: list[str] = []
    projects: list[Project] = []
    certifications: list[str] = []
    languages: list[str] = []


# ── Prompt injection protection ────────────────────────────────────────────────

# Patterns commonly used in prompt injection attacks.
# These phrases attempt to override the system prompt or inject new instructions.
_INJECTION_PATTERNS = re.compile(
    r"(ignore\s+(previous|all|above|prior)\s+(instructions?|prompts?|context)|"
    r"system\s*:|assistant\s*:|<\s*/?system\s*>|<\s*/?assistant\s*>|"
    r"\[INST\]|\[/INST\]|<\|im_start\|>|<\|im_end\|>|"
    r"###\s*(Instruction|System|Human|Assistant)|"
    r"you\s+are\s+now|pretend\s+you\s+are|act\s+as\s+if|"
    r"disregard\s+(your|the|all)|forget\s+(everything|all|your)|"
    r"new\s+instructions?\s*:)",
    re.IGNORECASE | re.MULTILINE,
)


def sanitize_for_ai(text: str, source: str = "input") -> str:
    """Sanitize user-supplied text before injecting it into an AI prompt.

    Defenses applied:
      1. Length truncation — limits tokens sent to Claude and prevents
         cost explosion from adversarially large payloads.
      2. Injection pattern removal — strips phrases commonly used to override
         system prompts (e.g. "Ignore previous instructions", "SYSTEM:").
      3. Null-byte removal — some LLM APIs behave unexpectedly with null bytes.

    This is defense-in-depth: the system prompt framing is the primary
    protection; this sanitizer is a secondary layer.

    Args:
        text:   The raw text to sanitize.
        source: A label for log messages (e.g. "resume_text", "career_goal").

    Returns:
        Sanitized text, safe to embed in an AI prompt.
    """
    if not text:
        return ""

    # 1. Truncate to configured limit
    max_len = security_settings.MAX_TEXT_LENGTH_FOR_AI
    if len(text) > max_len:
        logger.warning(
            "sanitize_for_ai: truncating %s from %d to %d chars (prompt injection / cost protection)",
            source,
            len(text),
            max_len,
        )
        text = text[:max_len]

    # 2. Remove null bytes (can cause unexpected behavior in LLM APIs)
    text = text.replace("\x00", "")

    # 3. Strip injection pattern matches (replace with a safe placeholder)
    cleaned, n_subs = _INJECTION_PATTERNS.subn("[redacted]", text)
    if n_subs > 0:
        logger.warning(
            "sanitize_for_ai: removed %d potential injection pattern(s) from %s",
            n_subs,
            source,
        )
        text = cleaned

    return text


# ── System prompt ───────────────────────────────────────────────────────────────

_SYSTEM_PROMPT = """You are an expert resume parser. Your job is to extract structured information from raw resume text and return it as valid JSON.

Extract the following fields from the resume text provided:
- full_name: The candidate's full name
- email: Email address
- phone: Phone number (as a string)
- location: City, state/country
- linkedin_url: LinkedIn profile URL
- github_url: GitHub profile URL
- portfolio_url: Personal website or portfolio URL
- summary: Professional summary or objective statement
- work_experience: Array of work history entries, each with:
  - company: Company/organization name
  - title: Job title/role
  - start_date: Start date (e.g., "Jan 2020", "2020-01", or "2020")
  - end_date: End date or "Present" if current
  - location: City/state (optional)
  - description: Array of bullet points/responsibilities (each as a separate string)
- education: Array of education entries, each with:
  - institution: School/university name
  - degree: Degree type (e.g., "Bachelor of Science", "MBA")
  - field: Field of study or major
  - graduation_year: Year of graduation or expected graduation
  - gpa: GPA if listed (as string)
- skills: Flat array of individual skill strings (e.g., ["Python", "React", "AWS"])
- projects: Array of project entries, each with:
  - name: Project name
  - description: Brief description
  - technologies: Array of technologies used
  - url: Project URL if mentioned
- certifications: Array of certification name strings
- languages: Array of spoken/written language strings (not programming languages)

Rules:
1. Return ONLY valid JSON — no markdown, no code fences, no explanations
2. Use null for missing fields, empty arrays [] for missing lists
3. Normalize dates to a consistent format where possible
4. Split combined skill strings into individual items (e.g., "Python, JavaScript" → ["Python", "JavaScript"])
5. Keep description bullet points concise but complete
6. If a URL is partial (e.g., "linkedin.com/in/johndoe"), expand it to the full URL (e.g., "https://linkedin.com/in/johndoe")
7. Do not invent information — only extract what is present in the text"""


# ── Text extraction ─────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str | Path | bytes) -> str:
    """Extract text from a PDF file using pdfplumber.

    Accepts a file path (str/Path) or raw bytes.
    Returns the combined text from all pages.
    """
    try:
        if isinstance(file_path, bytes):
            source = BytesIO(file_path)
        else:
            source = str(file_path)

        pages_text: list[str] = []
        with pdfplumber.open(source) as pdf:
            for i, page in enumerate(pdf.pages):
                try:
                    text = page.extract_text() or ""
                    pages_text.append(text)
                except Exception as exc:
                    logger.warning("Failed to extract text from PDF page %d: %s", i + 1, exc)
                    pages_text.append("")

        combined = "\n\n".join(t for t in pages_text if t.strip())
        if not combined.strip():
            raise TextExtractionError("PDF contains no extractable text (may be image-based)")
        return combined

    except TextExtractionError:
        raise
    except Exception as exc:
        raise TextExtractionError(f"Failed to extract text from PDF: {exc}") from exc


def extract_text_from_docx(file_path: str | Path | bytes) -> str:
    """Extract text from a DOCX file using python-docx.

    Accepts a file path (str/Path) or raw bytes.
    Returns the combined paragraph text.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise TextExtractionError("python-docx is not installed") from exc

    try:
        if isinstance(file_path, bytes):
            source = BytesIO(file_path)
        else:
            source = str(file_path)

        doc = Document(source)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        combined = "\n".join(paragraphs)
        if not combined.strip():
            raise TextExtractionError("DOCX contains no extractable text")
        return combined

    except TextExtractionError:
        raise
    except Exception as exc:
        raise TextExtractionError(f"Failed to extract text from DOCX: {exc}") from exc


# ── Claude parsing ──────────────────────────────────────────────────────────────

def parse_resume_with_claude(raw_text: str) -> ResumeData:
    """Parse resume text using Claude and return structured ResumeData.

    Uses prompt caching on the system prompt for cost efficiency.
    Retries up to 3 times with exponential backoff (2s, 4s, 8s).
    Raises ResumeParseError on failure after all retries.
    """
    client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)

    last_exc: Exception | None = None

    # Sanitize before sending to Claude — prevents prompt injection via resume content
    safe_text = sanitize_for_ai(raw_text, source="resume_text")

    for attempt in range(3):
        try:
            response = client.messages.create(
                model="claude-3-5-haiku-latest",
                max_tokens=4096,
                system=[
                    {
                        "type": "text",
                        "text": _SYSTEM_PROMPT,
                        "cache_control": {"type": "ephemeral"},
                    }
                ],
                messages=[
                    {
                        "role": "user",
                        "content": f"Parse this resume and return JSON:\n\n{safe_text}",
                    }
                ],
            )

            response_text = response.content[0].text.strip()

            # Strip markdown code fences if model wraps output anyway
            if response_text.startswith("```"):
                lines = response_text.split("\n")
                response_text = "\n".join(lines[1:-1] if lines[-1] == "```" else lines[1:])

            data = json.loads(response_text)
            return ResumeData(**data)

        except (json.JSONDecodeError, ValidationError) as exc:
            last_exc = exc
            logger.warning("Resume parse attempt %d — invalid response: %s", attempt + 1, exc)
        except anthropic.APIError as exc:
            last_exc = exc
            logger.warning("Resume parse attempt %d — API error: %s", attempt + 1, exc)
        except Exception as exc:
            last_exc = exc
            logger.warning("Resume parse attempt %d — unexpected error: %s", attempt + 1, exc)

        if attempt < 2:
            sleep_secs = 2 ** (attempt + 1)  # 2, 4, 8
            time.sleep(sleep_secs)

    raise ResumeParseError(
        f"Failed to parse resume after 3 attempts: {last_exc}"
    ) from last_exc


def parse_resume_bytes(file_bytes: bytes, file_ext: str) -> "ResumeData":
    """Extract text from raw bytes and parse with Claude. Used for inline (non-Celery) parsing."""
    if file_ext == "pdf":
        raw_text = extract_text_from_pdf(file_bytes)
    elif file_ext in ("docx", "doc"):
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raw_text = file_bytes.decode("utf-8", errors="ignore")
    return parse_resume_with_claude(raw_text)
